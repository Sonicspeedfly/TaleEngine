"""
Подготовка вложений-документов к отправке в нейросеть.

Нейросети по API не понимают .docx напрямую — поэтому:
  * PDF        -> отдаём как документ (Gemini читает PDF нативно);
  * DOCX/DOC/ODT/RTF -> если в системе есть LibreOffice (soffice) — конвертируем
                  в PDF (полная точность, включая картинки и вёрстку); иначе
                  извлекаем текст (python-docx для .docx) и отправляем помеченным
                  текстом — модель всё равно прочитает содержимое в один ход;
  * TXT/MD/CSV -> извлекаем текст.

Всё «по возможности»: если ни конвертация, ни извлечение не удались, кладём
короткую заметку, чтобы пользователь видел, что файл дошёл, но не распознан.

Формат PDF-блока — `image_url` с data:URI (как и картинки): LiteLLM определяет
mime по data:URI и шлёт Gemini как inline_data. Это тот же проверенный путь, что
и для изображений, поэтому работает на любой версии прокси.
"""
import base64
import io
import os
import re
import shutil
import subprocess
import tempfile

# Что считаем документом (не картинка/не аудио).
DOC_EXTS = {
    ".pdf", ".doc", ".docx", ".odt", ".rtf",
    ".txt", ".md", ".markdown", ".csv", ".log",
}
DOC_MIMES = {
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.oasis.opendocument.text",
    "application/rtf", "text/rtf",
    "text/plain", "text/markdown", "text/csv",
}
# Форматы, которые LibreOffice умеет открыть и сконвертировать в PDF.
_CONVERTIBLE = {".doc", ".docx", ".odt", ".rtf"}


def is_document(mime: str | None, name: str | None) -> bool:
    """Похоже ли вложение на документ (по mime или расширению имени файла)."""
    if mime and mime.lower().split(";")[0].strip() in DOC_MIMES:
        return True
    if name:
        return os.path.splitext(name)[1].lower() in DOC_EXTS
    return False


def _decode(data: str) -> bytes:
    """data может быть data:URI или голым base64."""
    if data.strip().lower().startswith("data:") and "," in data:
        data = data.split(",", 1)[1]
    return base64.b64decode(data)


def _ext_of(mime: str | None, name: str | None) -> str:
    if name and "." in os.path.basename(name):
        return os.path.splitext(name)[1].lower()
    m = (mime or "").lower()
    if "pdf" in m:
        return ".pdf"
    if "wordprocessingml" in m or "msword" in m:
        return ".docx"
    if "opendocument.text" in m:
        return ".odt"
    if "rtf" in m:
        return ".rtf"
    if "markdown" in m:
        return ".md"
    return ".txt"


def find_libreoffice() -> str | None:
    """Путь к soffice/libreoffice, если установлен (для конвертации в PDF)."""
    for exe in ("soffice", "libreoffice"):
        found = shutil.which(exe)
        if found:
            return found
    for path in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/soffice", "/usr/bin/libreoffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ):
        if os.path.exists(path):
            return path
    return None


def convert_to_pdf(src_bytes: bytes, ext: str) -> bytes | None:
    """Конвертировать документ в PDF через LibreOffice (если он есть). Иначе None."""
    exe = find_libreoffice()
    if not exe:
        return None
    try:
        with tempfile.TemporaryDirectory() as d:
            src = os.path.join(d, "input" + ext)
            with open(src, "wb") as f:
                f.write(src_bytes)
            subprocess.run(
                [exe, "--headless", "--convert-to", "pdf", "--outdir", d, src],
                check=True, timeout=90,
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
            )
            out = os.path.join(d, "input.pdf")
            if os.path.exists(out):
                with open(out, "rb") as f:
                    return f.read()
    except Exception:  # noqa: BLE001 — любая ошибка конвертации = откат на текст
        return None
    return None


def extract_docx_text(src_bytes: bytes) -> str:
    """Текст из .docx: абзацы (заголовки помечаем ##), плюс таблицы как 'a | b'."""
    try:
        import docx  # python-docx
    except ImportError:
        return ""
    try:
        doc = docx.Document(io.BytesIO(src_bytes))
    except Exception:  # noqa: BLE001
        return ""
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        style = (p.style.name if p.style else "") or ""
        parts.append(("## " + t) if "heading" in style.lower() else t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _decode_text(src_bytes: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1251", "latin-1"):
        try:
            return src_bytes.decode(enc).strip()
        except Exception:  # noqa: BLE001
            continue
    return ""


def _looks_binary(raw: bytes) -> bool:
    """
    Бинарный ли это файл (видео/архив/exe и т.п.), а не текст. Важно: latin-1
    «успешно» декодирует ЛЮБЫЕ байты, поэтому без этой проверки бинарник
    превращался в мегабайты мусора и ломал запрос к нейросети.
    """
    head = raw[:4096]
    if not head:
        return False
    if b"\x00" in head:
        return True
    try:
        head.decode("utf-8")
        return False
    except UnicodeDecodeError:
        pass
    # Не UTF-8: оцениваем долю непечатаемых символов в cp1251-прочтении.
    text = head.decode("cp1251", errors="replace")
    weird = sum(1 for ch in text if not (ch.isprintable() or ch in "\r\n\t"))
    return weird > len(text) * 0.15


def _pdf_block(pdf_bytes: bytes) -> dict:
    data_uri = "data:application/pdf;base64," + base64.b64encode(pdf_bytes).decode()
    return {"type": "image_url", "image_url": {"url": data_uri}}


def prepare_document(data: str, mime: str | None, name: str | None) -> dict:
    """
    Готовит content-блок для LLM из документа.
    Возвращает либо PDF-блок (image_url с data:application/pdf), либо текстовый
    блок {'type':'text','text': ...} с пометкой об имени файла.
    """
    raw = _decode(data)
    ext = _ext_of(mime, name)
    label = name or "документ"

    # PDF — как есть.
    if ext == ".pdf":
        if data.strip().lower().startswith("data:application/pdf"):
            return {"type": "image_url", "image_url": {"url": data}}
        return _pdf_block(raw)

    # Word/ODT/RTF — сперва пробуем настоящий PDF, иначе текст.
    if ext in _CONVERTIBLE:
        pdf = convert_to_pdf(raw, ext)
        if pdf:
            return _pdf_block(pdf)
        text = extract_docx_text(raw) if ext == ".docx" else _decode_text(raw)
        if text:
            return {"type": "text", "text": f"[Документ «{label}» — содержимое ниже]\n\n{text}"}
        return {"type": "text", "text": f"[Прикреплён документ «{label}», распознать не удалось.]"}

    # Бинарный файл неизвестного формата (видео/архив/…): мусор в контекст не льём.
    if _looks_binary(raw):
        return {
            "type": "text",
            "text": f"[Прикреплён файл «{label}» ({mime or 'неизвестный тип'}) — "
                    "формат не поддерживается для чтения, содержимое не передано.]",
        }

    # Текстовые форматы.
    text = _decode_text(raw)
    if text:
        return {"type": "text", "text": f"[Документ «{label}»]\n\n{text}"}
    return {"type": "text", "text": f"[Прикреплён документ «{label}».]"}


# ============================ ЭКСПОРТ КАНВАСА (Markdown -> Docx/PDF) ============================
def _find_unicode_font() -> str | None:
    """Системный TTF с кириллицей — для PDF (fpdf2 не умеет кириллицу core-шрифтами)."""
    for p in (
        r"C:\Windows\Fonts\arial.ttf", r"C:\Windows\Fonts\segoeui.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/Library/Fonts/Arial.ttf", "/System/Library/Fonts/Supplemental/Arial.ttf",
    ):
        if os.path.exists(p):
            return p
    return None


def _add_inline_runs(paragraph, text: str) -> None:
    """Текст с инлайн-разметкой **жирный**/*курсив*/`код` как runs (для .docx)."""
    for part in re.split(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def markdown_to_docx(content: str, title: str = "") -> bytes:
    """Markdown -> .docx: заголовки, списки, цитаты, код-блоки, инлайн-разметка."""
    import docx
    from docx.shared import Pt

    doc = docx.Document()
    if title:
        doc.add_heading(title, level=0)
    lines = (content or "").split("\n")
    i = 0
    while i < len(lines):
        st = lines[i].strip()
        if st.startswith("```"):
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            run = doc.add_paragraph().add_run("\n".join(code))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            i += 1
            continue
        h = re.match(r"^(#{1,6})\s+(.*)$", st)
        if h:
            doc.add_heading(h.group(2), level=min(len(h.group(1)), 4))
        elif re.match(r"^[-*]\s+", st):
            _add_inline_runs(doc.add_paragraph(style="List Bullet"), re.sub(r"^[-*]\s+", "", st))
        elif re.match(r"^\d+\.\s+", st):
            _add_inline_runs(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\.\s+", "", st))
        elif st.startswith(">"):
            doc.add_paragraph().add_run(st.lstrip("> ")).italic = True
        elif st == "":
            doc.add_paragraph()
        else:
            _add_inline_runs(doc.add_paragraph(), lines[i])
        i += 1
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def markdown_to_pdf(content: str, title: str = "") -> bytes:
    """
    Markdown -> PDF. Сначала пробуем точный путь docx->PDF через LibreOffice; если
    его нет — лёгкий PDF через fpdf2 с системным шрифтом (поддержка кириллицы).
    """
    try:
        pdf = convert_to_pdf(markdown_to_docx(content, title), ".docx")
        if pdf:
            return pdf
    except Exception:  # noqa: BLE001
        pass

    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    font = _find_unicode_font()
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    base = "Helvetica"
    if font:
        pdf.add_font("U", "", font)
        base = "U"

    def write(text: str, size: int = 11) -> None:
        pdf.set_font(base, "", size)
        # Без unicode-шрифта (редко) латинизируем, чтобы не упасть на кириллице.
        safe = text if base == "U" else text.encode("latin-1", "replace").decode("latin-1")
        # new_x/new_y возвращают курсор к левому полю на новую строку (иначе fpdf2
        # ругается «Not enough horizontal space», если предыдущая ячейка сместила X).
        pdf.multi_cell(0, 6, safe or " ", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    if title:
        write(title, 16)
        pdf.ln(2)
    in_code = False
    code: list[str] = []
    for line in (content or "").split("\n"):
        st = line.strip()
        if st.startswith("```"):
            if in_code:
                write("\n".join(code), 9)
                code = []
                pdf.ln(1)
            in_code = not in_code
            continue
        if in_code:
            code.append(line)
            continue
        h = re.match(r"^(#{1,6})\s+(.*)$", st)
        if h:
            write(h.group(2), 16 - len(h.group(1)))
            pdf.ln(1)
        elif st == "":
            pdf.ln(3)
        else:
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", st)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            clean = re.sub(r"`(.+?)`", r"\1", clean)
            write(clean, 11)
    return bytes(pdf.output())
